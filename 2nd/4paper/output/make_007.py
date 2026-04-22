#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_summary_document():
    """이온성 측사슬을 가진 빗질 고분자 논문 요약 docx 생성"""

    doc = Document()

    # 페이지 여백 설정 (2cm = 0.787inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.787)
        section.bottom_margin = Inches(0.787)
        section.left_margin = Inches(0.787)
        section.right_margin = Inches(0.787)

    # 논문 제목
    title = doc.add_paragraph()
    title_run = title.add_run("이온성 측사슬을 가진 빗질 고분자를 신규 시멘트 슬러리 분산제로의 활용: 합성, 특성화 및 작용 메커니즘")
    title_run.font.name = "맑은 고딕"
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.15

    # 연구 목적
    purpose_heading = doc.add_paragraph()
    purpose_heading_run = purpose_heading.add_run("연구 목적")
    purpose_heading_run.font.name = "맑은 고딕"
    purpose_heading_run.font.size = Pt(11)
    purpose_heading_run.font.bold = True
    purpose_heading.paragraph_format.space_before = Pt(6)
    purpose_heading.paragraph_format.space_after = Pt(3)
    purpose_heading.paragraph_format.line_spacing = 1.15

    purpose_text = doc.add_paragraph(
        "시멘트 슬러리의 유동성 유지 및 분산 안정성 향상을 위해 이온성 측사슬(ISC)을 도입한 신규 폴리카르복실산염 분산제(ISC-PCE)를 설계 및 합성하였다. "
        "이 연구는 기존의 PEG 기반 분산제의 한계를 극복하고, 혁신적인 분자 설계를 통해 우수한 콘크리트 유동성과 점탄성 특성을 제공하는 새로운 분산제 개발을 목표로 한다. "
        "특히 이온성 측사슬의 도입이 시멘트 입자에 대한 흡착 및 안정화 메커니즘을 규명하는 데 중점을 둔다."
    )
    purpose_text.paragraph_format.space_after = Pt(3)
    purpose_text.paragraph_format.line_spacing = 1.15
    for run in purpose_text.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)

    # 주요 방법
    method_heading = doc.add_paragraph()
    method_heading_run = method_heading.add_run("주요 방법")
    method_heading_run.font.name = "맑은 고딕"
    method_heading_run.font.size = Pt(11)
    method_heading_run.font.bold = True
    method_heading.paragraph_format.space_before = Pt(6)
    method_heading.paragraph_format.space_after = Pt(3)
    method_heading.paragraph_format.line_spacing = 1.15

    methods = [
        "2단계 간단한 합성법을 통해 이온성 측사슬 폴리카르복실산염(ISC-PCE) 분산제 제조",
        "FT-IR, 1H NMR, 크기 배제 크로마토그래피(SEC) 등을 이용한 분자 구조 특성화",
        "유동성 유지율, 자유 침강 시간, 미세 형태 분석을 통한 성능 평가 및 분산 안정성 검증",
        "분자 배치, 흡착량, 제타 전위, 복합화 능력, 흡착층 두께 측정을 통한 작용 메커니즘 규명"
    ]

    for method in methods:
        p = doc.add_paragraph(method, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = "맑은 고딕"
            run.font.size = Pt(10)

    # 핵심 결과
    result_heading = doc.add_paragraph()
    result_heading_run = result_heading.add_run("핵심 결과")
    result_heading_run.font.name = "맑은 고딕"
    result_heading_run.font.size = Pt(11)
    result_heading_run.font.bold = True
    result_heading.paragraph_format.space_before = Pt(6)
    result_heading.paragraph_format.space_after = Pt(3)
    result_heading.paragraph_format.line_spacing = 1.15

    results = [
        "ISC-PCE 함유 시멘트 페이스트는 3시간 후 약 79.6%의 우수한 유동성 유지율을 달성",
        "3시간 이상의 긴 자유 침강 시간과 감소된 응집 구조로 우수한 분산 안정성 입증",
        "이온성 측사슬 도입으로 인해 기존 PCE 대비 흡착층 두께가 103.1% 증가",
        "부드럽고 연속적인 흡착 및 완화된 복합화 능력을 통한 개선된 작용 메커니즘 규명"
    ]

    for result in results:
        p = doc.add_paragraph(result, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = "맑은 고딕"
            run.font.size = Pt(10)

    # 결론
    conclusion_heading = doc.add_paragraph()
    conclusion_heading_run = conclusion_heading.add_run("결론")
    conclusion_heading_run.font.name = "맑은 고딕"
    conclusion_heading_run.font.size = Pt(11)
    conclusion_heading_run.font.bold = True
    conclusion_heading.paragraph_format.space_before = Pt(6)
    conclusion_heading.paragraph_format.space_after = Pt(3)
    conclusion_heading.paragraph_format.line_spacing = 1.15

    conclusion_text = doc.add_paragraph(
        "이온성 측사슬을 도입한 ISC-PCE는 기존의 비이온성 PCE를 능가하는 성능을 보이며, "
        "특히 높은 슬럼프 유지 요구사항이 있는 시멘트 시스템에서 효과적인 대체 분산제로 활용될 수 있다. "
        "본 연구는 새로운 측사슬 단량체 개발의 방향성을 제시하고, 분자 설계를 통한 시멘트 기반 재료의 성능 개선에 대한 이론적 기초를 제공한다."
    )
    conclusion_text.paragraph_format.space_after = Pt(0)
    conclusion_text.paragraph_format.line_spacing = 1.15
    for run in conclusion_text.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)

    # 문서 저장
    output_path = r"E:\venture\proposal\JBLab\2nd\4paper\output\007.docx"
    doc.save(output_path)

    print(f"Success: Document created at {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        create_summary_document()
        print("007.docx 파일이 성공적으로 생성되었습니다.")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

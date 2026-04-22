#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 한글 폰트 설정
FONT_NAME = "맑은 고딕"
FONT_NAME_ENG = "Calibri"

# Document 생성
doc = Document()

# 여백 설정 (2cm)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# 제목
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("폴리머 초소성제, 방수제 및 포졸라나 물질 조합을\n이용한 역사적 석조 보수용 에어 라임 그라우트 개선")
title_run.font.name = FONT_NAME
title_run.font.size = Pt(14)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)

# 줄간격 설정
title.paragraph_format.line_spacing = 1.15
title.paragraph_format.space_after = Pt(6)

# 구분선
separator1 = doc.add_paragraph()
separator1.paragraph_format.line_spacing = 1.15
separator1.paragraph_format.space_after = Pt(3)

# 연구 목적
heading1 = doc.add_paragraph()
heading1_run = heading1.add_run("연구 목적")
heading1_run.font.name = FONT_NAME
heading1_run.font.size = Pt(11)
heading1_run.font.bold = True
heading1.paragraph_format.line_spacing = 1.15
heading1.paragraph_format.space_after = Pt(2)

content1 = doc.add_paragraph()
content1_text = ("역사적 석조 건축물의 보수는 원본 자재와의 호환성을 유지하면서 내구성과 성능을 향상시켜야 한다. "
                "본 연구는 폴리머 초소성제, 방수제, 포졸라나 물질을 조합하여 에어 라임 기반 그라우트의 성능을 개선하는 방법을 제시한다. "
                "이를 통해 역사 건축물 보존에 적합한 고성능 에코프렌들리 그라우트 개발을 목표로 한다.")
content1_run = content1.add_run(content1_text)
content1_run.font.name = FONT_NAME
content1_run.font.size = Pt(10)
content1.paragraph_format.line_spacing = 1.15
content1.paragraph_format.space_after = Pt(4)

# 주요 방법
heading2 = doc.add_paragraph()
heading2_run = heading2.add_run("주요 방법")
heading2_run.font.name = FONT_NAME
heading2_run.font.size = Pt(11)
heading2_run.font.bold = True
heading2.paragraph_format.line_spacing = 1.15
heading2.paragraph_format.space_after = Pt(2)

methods = [
    "폴리카복실산계 초소성제의 첨가로 그라우트의 유동성 및 작업성 개선",
    "소수성 물질(방수제)을 혼입하여 장기 내구성 및 내수성 향상",
    "포졸라나 물질의 첨가로 에어 라임의 강도 및 내구성 증진",
    "다양한 혼합 비율의 조건에서 압축강도, 투과성, 공극 구조 등의 물성 평가",
    "역사 건축물 원래 자재와의 호환성 및 미학적 특성 검증"
]

for method in methods:
    p = doc.add_paragraph(method, style='List Bullet')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.15
    p_format.space_after = Pt(1)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(10)

# 줄간격 조정
doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

# 핵심 결과
heading3 = doc.add_paragraph()
heading3_run = heading3.add_run("핵심 결과")
heading3_run.font.name = FONT_NAME
heading3_run.font.size = Pt(11)
heading3_run.font.bold = True
heading3.paragraph_format.line_spacing = 1.15
heading3.paragraph_format.space_after = Pt(2)

results = [
    "초소성제 첨가로 그라우트의 유동성이 현저히 증가하여 시공성 개선",
    "방수제의 적용으로 물 흡수율 감소 및 장기 내구성 향상 확인",
    "포졸라나 물질 혼입 시 강도 발달이 양호하면서 에코프렌들리 특성 유지",
    "최적 배합 조건에서 역사 건축물 기준을 만족하는 성능 달성",
    "원본 석조체와 우수한 호환성으로 보존 가치 유지"
]

for result in results:
    p = doc.add_paragraph(result, style='List Bullet')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.15
    p_format.space_after = Pt(1)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(10)

# 줄간격 조정
doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

# 결론
heading4 = doc.add_paragraph()
heading4_run = heading4.add_run("결론")
heading4_run.font.name = FONT_NAME
heading4_run.font.size = Pt(11)
heading4_run.font.bold = True
heading4.paragraph_format.line_spacing = 1.15
heading4.paragraph_format.space_after = Pt(2)

conclusion = doc.add_paragraph()
conclusion_text = ("폴리머 초소성제, 방수제, 포졸라나 물질의 조합은 에어 라임 기반 그라우트의 성능을 현저히 개선할 수 있음을 입증했다. "
                  "이 연구 결과는 역사적 석조 건축물의 보수 및 보존에 있어 새로운 친환경적이고 실용적인 솔루션을 제공한다.")
conclusion_run = conclusion.add_run(conclusion_text)
conclusion_run.font.name = FONT_NAME
conclusion_run.font.size = Pt(10)
conclusion.paragraph_format.line_spacing = 1.15

# 문서 저장
output_path = "E:/venture/proposal/JBLab/2nd/4paper/output/008.docx"
doc.save(output_path)
print(f"Document saved successfully: {output_path}")

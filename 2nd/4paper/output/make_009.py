# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Set margins (2cm = 0.787 inches)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.787)
    section.bottom_margin = Inches(0.787)
    section.left_margin = Inches(0.787)
    section.right_margin = Inches(0.787)

# Set line spacing to 1.15 for the entire document
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.15

# Title
title = doc.add_paragraph()
title_run = title.add_run('하이브리드 측쇄를 가진 신규 블록 폴리카복실산염의 탄-수 슬러리에서의 분산 성능 및 흡착 거동')
title_run.font.name = '맑은 고딕'
title_run.font.size = Pt(14)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.paragraph_format
title_format.line_spacing = 1.15
title_format.space_after = Pt(6)

# Research Purpose
doc.add_paragraph('연구 목적', style='List Bullet').runs[0].font.bold = True
purpose_heading = doc.paragraphs[-1]
purpose_heading.style = None
purpose_run = purpose_heading.add_run('연구 목적')
purpose_run.font.name = '맑은 고딕'
purpose_run.font.size = Pt(11)
purpose_run.font.bold = True
purpose_heading.paragraph_format.line_spacing = 1.15
purpose_heading.paragraph_format.space_after = Pt(3)

purpose = doc.add_paragraph('저등급 석탄의 고농도 슬러리 제조를 위해 장-단 복합 측쇄를 가진 새로운 블록 폴리카복실산염(PCD) 분산제를 합성하고 특성을 조사하였다. 측쇄 구조와 탄-수 슬러리(CWS)의 성질 간의 관계를 규명하고, 분산제의 흡착 거동이 분산 성능에 미치는 영향을 이해하는 것을 목표로 하였다.')
purpose_run = purpose.runs[0]
purpose_run.font.name = '맑은 고딕'
purpose_run.font.size = Pt(10)
purpose.paragraph_format.line_spacing = 1.15
purpose.paragraph_format.space_after = Pt(3)

# Main Methods
methods_heading = doc.add_paragraph('주요 방법')
methods_run = methods_heading.runs[0]
methods_run.font.name = '맑은 고딕'
methods_run.font.size = Pt(11)
methods_run.font.bold = True
methods_heading.paragraph_format.line_spacing = 1.15
methods_heading.paragraph_format.space_after = Pt(3)

methods = [
    'RAFT 중합을 이용하여 다양한 단-장 측쇄 몰 비율을 가진 블록 폴리카복실산염 분산제 5종(PCD-1~PCD-5) 합성',
    '탄-수 슬러리의 겉보기 점도, 제타 전위, 안정성을 측정하여 분산 성능 평가',
    '총유기탄소 분석기를 이용하여 석탄 표면에 흡착된 분산제의 양 정량화',
    '주사전자현미경(ESEM)과 에너지 분산 X선 분광(EDS)으로 석탄 표면의 분산제 흡착 상태 관찰',
    '분산제 농도별 흡착 등온선 측정 및 레올로지 특성(Herschel-Bulkley 및 Power-law 모델) 분석'
]

for method in methods:
    m = doc.add_paragraph(method, style='List Bullet')
    m_run = m.runs[0]
    m_run.font.name = '맑은 고딕'
    m_run.font.size = Pt(10)
    m.paragraph_format.line_spacing = 1.15
    m.paragraph_format.space_after = Pt(2)
    m.paragraph_format.left_indent = Inches(0.25)

# Core Results
results_heading = doc.add_paragraph('핵심 결과')
results_run = results_heading.runs[0]
results_run.font.name = '맑은 고딕'
results_run.font.size = Pt(11)
results_run.font.bold = True
results_heading.paragraph_format.line_spacing = 1.15
results_heading.paragraph_format.space_after = Pt(3)

results = [
    '단:장 측쇄 몰 비율이 약 1:1인 PCD-2가 최적의 분산 성능과 안정성을 나타냈으며, RAFT 중합으로 제조한 PCD가 무작위 중합 제어 표본보다 우수한 점도 감소 효과 시현',
    '석탄 표면에 대한 PCD의 흡착량이 증가함에 따라 슬러리의 겉보기 점도, 안정성, 제타 전위가 모두 향상되었으므로 흡착 거동이 분산 성능의 핵심 인자임을 확인',
    'PCD-2는 석탄 표면에서 더 큰 흡착량을 형성하여 정전기적 반발력과 입체 장애 효과를 극대화함으로써 우수한 분산 특성 발휘',
    '고온 및 고전단 환경에서도 PCD의 흡착층이 안정적으로 유지되어 저등급 석탄의 장시간 저장 및 수송 시 슬러리 점도 증가를 억제'
]

for result in results:
    r = doc.add_paragraph(result, style='List Bullet')
    r_run = r.runs[0]
    r_run.font.name = '맑은 고딕'
    r_run.font.size = Pt(10)
    r.paragraph_format.line_spacing = 1.15
    r.paragraph_format.space_after = Pt(2)
    r.paragraph_format.left_indent = Inches(0.25)

# Conclusion
conclusion_heading = doc.add_paragraph('결론')
conclusion_run = conclusion_heading.runs[0]
conclusion_run.font.name = '맑은 고딕'
conclusion_run.font.size = Pt(11)
conclusion_run.font.bold = True
conclusion_heading.paragraph_format.line_spacing = 1.15
conclusion_heading.paragraph_format.space_after = Pt(3)

conclusion = doc.add_paragraph('하이브리드 측쇄를 가진 블록 폴리카복실산염 분산제는 장쇄의 공간적 장애 효과와 단쇄의 흡착 용이성을 결합하여 저등급 석탄 슬러리의 분산 성능을 현저히 개선한다. 특히 단:장 측쇄 비율의 최적화를 통해 석탄 표면에서의 흡착 거동을 정밀하게 제어할 수 있으며, 이는 고농도 탄-수 슬러리의 개발 및 실용화에 중요한 지침을 제공한다.')
conclusion_run = conclusion.runs[0]
conclusion_run.font.name = '맑은 고딕'
conclusion_run.font.size = Pt(10)
conclusion.paragraph_format.line_spacing = 1.15

# Save the document
output_path = 'E:/venture/proposal/JBLab/2nd/4paper/output/009.docx'
doc.save(output_path)
print(f"Document created successfully: {output_path}")
